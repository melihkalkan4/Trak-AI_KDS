from pathlib import Path
import shutil

from docx import Document


INPUT = Path(r"C:\Users\Melih Kalkan\Downloads\TRAK_AI_DSS_Thesis_FINAL (1).docx")
OUTPUT = Path(r"C:\Users\Melih Kalkan\Desktop\Trak-AI_KDS\TRAK_AI_DSS_Thesis_FINAL_references_corrected.docx")


def normalized(text: str) -> str:
    return " ".join(text.split())


def set_para_text(paragraph, text: str) -> None:
    paragraph.text = text


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)

    shutil.copy2(INPUT, OUTPUT)
    doc = Document(OUTPUT)

    replacements = {
        "CSIRO: Commonwealth Scientific and Industrial Research Organisation": "DAWE: Department of Agriculture, Water and the Environment",
        "Provisional Canonical References": "Additional References",
    }

    exact_reference_replacements = {
        "Abdelmoneim, A., Abdrabou, A., Said, M., & El-Sayed, R. (2025). Smart digital soil sensors for precision agriculture: A review. Sensors, 25(2), 343. https://doi.org/10.3390/s25020343":
            "Abdelmoneim, A. A., Al Kalaany, C. M., Khadra, R., Derardja, B., & Dragonetti, G. (2025). Calibration of low-cost capacitive soil moisture sensors for irrigation management applications. Sensors, 25(2), 343. https://doi.org/10.3390/s25020343",
        "Boursianis, A. D., Papadopoulou, M. S., Diamantoulakis, P., Liopa-Tsakalidi, A., Barouchas, P., Salahas, G., Karagiannidis, G., Wan, S., & Goudos, S. K. (2022). Internet of Things (IoT) and Agricultural Unmanned Aerial Vehicles (UAVs) in smart farming: A comprehensive review. Internet of Things, 18, 100187.":
            "Boursianis, A. D., Papadopoulou, M. S., Diamantoulakis, P., Liopa-Tsakalidi, A., Barouchas, P., Salahas, G., Karagiannidis, G. K., Wan, S., & Goudos, S. K. (2022). Internet of Things (IoT) and Agricultural Unmanned Aerial Vehicles (UAVs) in smart farming: A comprehensive review. Internet of Things, 18, 100187. https://doi.org/10.1016/j.iot.2020.100187",
        "Commonwealth Scientific and Industrial Research Organisation (CSIRO). (2023). National artificial intelligence in agriculture strategy. Australian Government Department of Agriculture, Fisheries and Forestry.":
            "Department of Agriculture, Water and the Environment. (2022). Digital Foundations for Agriculture Strategy: Driving the development and uptake of digital technologies in the Australian agriculture, fisheries and forestry industry. Australian Government. https://www.agriculture.gov.au/sites/default/files/documents/digital-foundations-agriculture-strategy.pdf",
        "David, E., Madec, S., Sadeghi-Tehran, P., Aasen, H., Zheng, B., Liu, S., Kirchgessner, N., Ishikawa, G., Nagasawa, K., Badhon, M. A., Pozniak, C., Solan, B., Hund, A., Chapman, S. C., Baret, F., Stavness, I., & Guo, W. (2021). Global Wheat Head Detection 2021: An improved dataset for benchmarking wheat head detection methods. Plant Phenomics, 2021, 9846158.":
            "David, E., Serouart, M., Smith, D., Madec, S., Velumani, K., Liu, S., Wang, X., Pinto, F., Shafiee, S., Tahir, I. S. A., Tsujimoto, H., Nasuda, S., Zheng, B., Kirchgessner, N., Aasen, H., Hund, A., Sadeghi-Tehran, P., Nagasawa, K., Ishikawa, G., ... Guo, W. (2021). Global Wheat Head Detection 2021: An improved dataset for benchmarking wheat head detection methods. Plant Phenomics, 2021, 9846158. https://doi.org/10.34133/2021/9846158",
        "Meier, U. (Ed.). (2018). Growth stages of mono- and dicotyledonous plants: BBCH Monograph (2nd ed.). Open Agrar Repositorium, Julius Kühn-Institut.":
            "Meier, U. (Ed.). (2018). Growth stages of mono- and dicotyledonous plants: BBCH Monograph (2nd ed.). Open Agrar Repositorium, Julius Kühn-Institut. https://doi.org/10.5073/20180906-074619",
        "Miller, J., Patel, R., Hou, S., Ramirez, D., Anderson, K., Brown, T., Wilson, M., & Stewart, L. (2025). Rural infrastructure constraints and the next generation of smart farming systems. Agricultural Systems, 217, 103938.":
            "Göre, M. E. (2009). Epidemic outbreaks of downy mildew caused by Plasmopara halstedii on sunflower in Thrace, part of the Marmara region of Turkey. Plant Pathology, 58(2), 396.",
        "Nagahage, E. A. A. D., Nagahage, I. S. P., & Fujino, T. (2019). Calibration and validation of a low-cost capacitive moisture sensor to integrate the automated soil moisture monitoring system. Agriculture, 9(7), 141.":
            "Nagahage, E. A. A. D., Nagahage, I. S. P., & Fujino, T. (2019). Calibration and validation of a low-cost capacitive moisture sensor to integrate the automated soil moisture monitoring system. Agriculture, 9(7), 141. https://doi.org/10.3390/agriculture9070141",
        "Scientific and Technological Research Council of Türkiye (TÜBİTAK). (2024). 2209/A Undergraduate Students Research Projects Funding Programme. TÜBİTAK ARDEB.":
            "Scientific and Technological Research Council of Türkiye (TÜBİTAK). (2024). 2209-A Undergraduate Students Research Projects Funding Programme. TÜBİTAK BİDEB. https://tubitak.gov.tr/en/scholarships/degree-associate-degree/destek-programlari/2209-research-project-support-programme-undergraduate-students",
        "Tzachor, A., Devare, M., Richards, C., Pypers, P., Ghosh, A., Koo, J., Johal, S., & King, B. (2023). Large language models and agricultural extension services. Nature Food, 4, 941–948.":
            "Tzachor, A., Devare, M., Richards, C., Pypers, P., Ghosh, A., Koo, J., Johal, S., & King, B. (2023). Large language models and agricultural extension services. Nature Food, 4, 941–948. https://doi.org/10.1038/s43016-023-00867-x",
        "Bogena, H. R., Huisman, J. A., Schilling, B., Güntner, A., & Vereecken, H. (2017). Effective calibration of low-cost soil water content sensors. Sensors, 17(1), 208.":
            "Bogena, H. R., Huisman, J. A., Schilling, B., Weuthen, A., & Vereecken, H. (2017). Effective calibration of low-cost soil water content sensors. Sensors, 17(1), 208. https://doi.org/10.3390/s17010208",
        "Friedrich, T. (2017). A new paradigm for feeding the world in 2050: The sustainable intensification of crop production. Resource Magazine, 24(2), 18.":
            "Friedrich, T. (2015). A new paradigm for feeding the world in 2050: The sustainable intensification of crop production. Resource Magazine, 22(2), 18.",
        "Liu, J., Shen, D., Zhang, Y., Dolan, B., Carin, L., & Chen, W. (2023). What makes good in-context examples for GPT-3? In Proceedings of the 1st Workshop on Deep Learning Inside Out (DeeLIO).":
            "Liu, J., Shen, D., Zhang, Y., Dolan, B., Carin, L., & Chen, W. (2022). What makes good in-context examples for GPT-3? In Proceedings of the 3rd Workshop on Deep Learning Inside Out (DeeLIO 2022) (pp. 100–114). Association for Computational Linguistics. https://doi.org/10.18653/v1/2022.deelio-1.10",
        "Picon, A., Seitz, M., Alvarez-Gila, A., Mohnke, P., Ortiz-Barredo, A., & Echazarra, J. (2022). Crop conditional convolutional neural networks for massive multi-crop plant disease classification over cell phone acquired images. Computers and Electronics in Agriculture, 167, 105093.":
            "Picon, A., Seitz, M., Alvarez-Gila, A., Mohnke, P., Ortiz-Barredo, A., & Echazarra, J. (2019). Crop conditional convolutional neural networks for massive multi-crop plant disease classification over cell phone acquired images taken on real field conditions. Computers and Electronics in Agriculture, 167, 105093. https://doi.org/10.1016/j.compag.2019.105093",
        "Schaefer, M. T., & Lamm, R. D. (2017). A review of soil moisture sensors and their use in precision irrigation. Agricultural Water Management, 184, 9–20.":
            "Hardie, M. (2020). Review of novel and emerging proximal soil moisture sensors for use in agriculture. Sensors, 20(23), 6934. https://doi.org/10.3390/s20236934",
        "Bangladesh Agricultural Research Institute (BARI). (2022). Sunflower disease dataset. BARI Plant Pathology Division.":
            "Göçmen, E. (2020). The effect of different irrigation treatment on yield and yield parameters of sunflowers in semi-arid conditions. Journal of Tekirdag Agricultural Faculty, 18(1), 80–90. https://doi.org/10.33462/jotaf.716280",
        "Edirne Provincial Directorate of Agriculture and Forestry. (2025). 2024 activity report and 2025 surveillance summary. Republic of Türkiye Ministry of Agriculture and Forestry, Edirne Provincial Directorate.":
            "Edirne İl Tarım ve Orman Müdürlüğü. (2025). 2025 yılı Edirne İl Tarım ve Orman Müdürlüğü faaliyet raporu. T.C. Tarım ve Orman Bakanlığı. https://edirne.tarimorman.gov.tr/Belgeler/Edirne%20%C4%B0l%20Tar%C4%B1m%20ve%20Orman%20M%C3%BCd%C3%BCrl%C3%BC%C4%9F%C3%BC%202025%20Y%C4%B1l%C4%B1%20Faaliyet%20Raporu.pdf",
        "General Directorate of Agricultural Research and Policies (TAGEM). (2024). Septoria leaf-spot prevalence and resistance breeding in Thracian wheat varieties (Project Final Report). Republic of Türkiye Ministry of Agriculture and Forestry.":
            "Tohumluk Tescil ve Sertifikasyon Merkez Müdürlüğü. (2019). Trakya Bölgesi ekmeklik buğday tescil raporu: Aleppo, Pandiya, Waximum, Anafarta (BBVD12-2016), Abide (BBVD17-2016) ve Stoyana. T.C. Tarım ve Orman Bakanlığı. https://www.tarimorman.gov.tr/BUGEM/TTSM/Belgeler/Yay%C4%B1nlar/Tescil%20Raporlar%C4%B1/2019/Ekmeklik%20Bu%C4%9Fday/trakyaekm19.pdf",
        "Indian Agricultural Research Institute (IARI). (2021). Wheat nitrogen deficiency and rust imagery dataset. IARI Division of Agricultural Physics.":
            "Oksal, E., & Maden, S. (2019). Determination of the races of Plasmopara halstedii (Farl.) Berl. & de Toni, the causal agent of sunflower downy mildew in Turkey and reactions of some commercial sunflower varieties against these races. Plant Protection Bulletin, 59(4), 77–84. https://doi.org/10.16955/bitkorb.635663",
        "Kaggle Wheat Disease Dataset Contributors. (2022). Wheat disease 21K dataset. Kaggle Platform.":
            "Öztürk, İ., Kahraman, T., Avcı, R., Girgin, V. Ç., Çiftçigil, T. H., Seidi, M., Tülek, A., Akın, K., & Tuna, B. (2018). Ekmeklik buğday (Triticum aestivum L.) genotiplerinde çevre koşullarının agronomik karakterler ve biyotik stres faktörlerine etkisi. Bahri Dağdaş Bitkisel Araştırma Dergisi, 7(1), 14–22. https://izlik.org/JA94CW54NB",
        "Thrace Agricultural Research Institute. (2024). Sunflower downy mildew and broomrape population dynamics in the Thracian basin (Technical Report). Republic of Türkiye Ministry of Agriculture and Forestry, TAGEM.":
            "Üder, F., & Demirbaş, S. (2019). Trakya bölgesi canavar otlarının (Orobanche cumana Wallr.) ayçiçeğinin gelişimi üzerine bazı etkilerinin belirlenmesi. Mediterranean Agricultural Sciences, 32(2), 211–217. https://doi.org/10.29136/mediterranean.567233",
        "Trakya University. (2023). Genotype × environment interaction in bread wheat across Edirne, Lüleburgaz, and Tekirdağ sub-zones. Trakya University Faculty of Agriculture Technical Report.":
            "Karabulut, A. A., Ceylan, N., Bahar, E., & Kurşun, İ. (2021). Crop phenology-based object-oriented classification approach using SENTINEL-2A and NDVI time series: Sunflower crops in Kırklareli, Turkey. International Journal of Environment and Geoinformatics, 8(3), 316–327. https://doi.org/10.30897/ijegeo.858456",
        "Trakya University Agricultural Faculty. (2023). Water–yield response of sunflower under regulated deficit irrigation in the Tekirdağ region. Trakya University Faculty of Agriculture Doctoral Studies Programme.":
            "Gürbüz, M. A., Kayalı, E., Bahar, E., Öz, T. A., & Kurşun, İ. (2019). Trakya topraklarının veri tabanının oluşturulması ve bazı toprak özellikleri. Toprak Bilimi ve Bitki Besleme Dergisi, 7(1), 28–36. https://doi.org/10.33409/tbbbd.595133",
        "Trakya University Journal of Natural Sciences. (2023). Common bunt and sun-pest (Süne) biological control practices in the Saray district of Tekirdağ. Trakya University Journal of Natural Sciences.":
            "Pekcan, V., & Erdem, T. (2016). Edirne koşullarında destekleme sulamanın ayçiçeğinin su kullanımı ve verimine etkileri. Trakya Üniversitesi Fen Bilimleri Dergisi, 6(2), 59–66. https://izlik.org/JA34MH39FU",
    }

    startswith_replacements = {
        "The first gap is the cost–accuracy dilemma.": (
            "The first gap is the cost–accuracy dilemma. Industrial-grade precision-agriculture solutions employ professional soil probes, calibrated weather stations, and proprietary cloud analytics whose per-hectare cost places them beyond the reach of typical Thracian wheat and sunflower operations. Conversely, hobbyist-grade IoT sensors, although affordable, are well documented to produce systematically biased readings that depend on local soil texture, salinity, and temperature (Nagahage et al., 2019) . In the absence of a credible calibration strategy, the inexpensive path yields decisions of unacceptable quality, whereas the accurate path remains economically inaccessible. Recent calibration evidence indicates that polynomial calibration of low-cost capacitive probes against gravimetric reference samples can restore usable accuracy at silty-clay sites (Abdelmoneim et al., 2025) ; a viable system for the Thrace region must therefore reconcile capital efficiency with scientifically defensible measurement fidelity."
        ),
        "The fifth gap is the absence of regionally localised models.": (
            "The fifth gap is the absence of regionally localised models. Global agricultural artificial-intelligence products are typically trained on broad, geographically heterogeneous datasets that average over micro-climates and soil profiles which, in practice, behave heterogeneously. Thrace exhibits distinctive characteristics, clay-loam to silty-clay soils, a continental–Mediterranean transitional climate, and rotation-based cropping practices, that are inadequately represented by such global baselines. The Australian Government's Digital Foundations for Agriculture Strategy reframes the issue methodologically: digital agriculture should be developed around the operating conditions, data foundations, and adoption constraints of the farming systems in which it is deployed (Department of Agriculture, Water and the Environment, 2022) . The motivation of the present thesis is to take that question seriously for the Thrace region (Edirne İl Tarım ve Orman Müdürlüğü, 2025) ."
        ),
        "Several authors have proposed taxonomies of the technological generations through which precision agriculture has progressed.": (
            "Several authors have proposed taxonomies of the technological generations through which precision agriculture has progressed. (Wolfert et al., 2017) describe an emergent phase that they characterise as smart farming, in which the locus of innovation shifts from individual machinery to integrated data platforms that synthesise sensor networks, machine-learning models, and cloud analytics. More recent literature speaks of Agriculture 4.0 and, increasingly, Agriculture 5.0, the latter emphasising human-centred artificial intelligence and the integration of autonomous robotics with explicit decision-making support (Klerkx et al., 2019) (Saiz-Rubio & Rovira-Más, 2020) . Across these generational labels, two recurring themes are persistently visible: the growing centrality of data as a productive factor of agricultural production, and the persistent difficulty of bringing precision-agriculture technologies into the hands of small and medium farmers who, in aggregate, manage the majority of the world's cultivated land (Schimmelpfennig, 2016) (Friedrich, 2015) ."
        ),
        "A critical but frequently under-appreciated challenge in agricultural IoT is sensor calibration.": (
            "A critical but frequently under-appreciated challenge in agricultural IoT is sensor calibration. Low-cost capacitive soil-moisture sensors such as the DFRobot SKU:SEN0193 produce raw analogue readings that depend strongly on soil texture, salinity, and temperature, the very factors that vary most across a heterogeneous landscape (Nagahage et al., 2019) (Hardie, 2020) . Without site-specific calibration, the readings produced by such sensors carry root-mean-square errors of approximately 0.08 m³/m³ volumetric water content, an error magnitude large enough to compromise irrigation decisions in marginal-water regimes. After polynomial calibration against gravimetric reference samples, the same sensors can attain errors below 0.03 m³/m³, restoring usable accuracy. (Abdelmoneim et al., 2025) report comparable improvements at silty-clay sites with coefficient-of-determination values between 0.85 and 0.87 after sensor-specific calibration, and (Bogena et al., 2017) document the calibration protocols underlying the SoilNet wireless sensor network deployments in Germany and the Netherlands. These convergent findings establish that the low-cost path to soil-moisture sensing is scientifically defensible provided that a credible calibration strategy is adopted; the absence of such a strategy is the primary source of skepticism toward consumer-grade IoT in the academic precision-agriculture literature, and motivates the polynomial-regression calibration procedure documented in Section 3.2 and Section 4.3 of the present work."
        ),
        "For phenological staging, the Biologische Bundesanstalt": (
            "For phenological staging, the Biologische Bundesanstalt, Bundessortenamt und Chemische Industrie (BBCH) scale provides a numerical encoding of plant developmental stages that is internationally standardised across crop species (Meier, 2018) . Recent computer-vision research has begun to treat BBCH stage estimation as a classification or regression problem, leveraging large labelled datasets such as the Global Wheat Head Detection 2021 corpus (David et al., 2021) , which contains 6,500 images and approximately 275,000 wheat-head annotations collected from sites distributed across Europe, Asia, North America, and other wheat-growing regions. The labelling of plant developmental stages remains, however, more subjective than disease labelling, and large open BBCH datasets specific to sunflower remain comparatively scarce, motivating the use of smaller curated sources such as the BARI Sunflower Disease dataset within the present work."
        ),
        "At the line-ministry level, the Republic of Türkiye Ministry of Agriculture and Forestry": (
            "At the line-ministry level, the Republic of Türkiye Ministry of Agriculture and Forestry, through its research and development arm, the General Directorate of Agricultural Research and Policies (TAGEM), funds and publishes a substantial body of regional research relevant to wheat, sunflower, irrigation, and disease management in the Thrace basin. Several of these publications are incorporated into the TRAK-AI DSS knowledge base, including the Edirne İl Tarım ve Orman Müdürlüğü 2025 activity report (Edirne İl Tarım ve Orman Müdürlüğü, 2025) and official Trakya-region wheat variety registration reports that include yield, quality, and disease-observation data (Tohumluk Tescil ve Sertifikasyon Merkez Müdürlüğü, 2019) . The Thrace Agricultural Research Institute, located in Edirne and operating under TAGEM, conducts continuous variety trials and disease surveillance that constitute an authoritative source of regionally calibrated agronomic guidance."
        ),
        "Studies focusing on bread wheat in the Thracian basin examine": (
            "Studies focusing on bread wheat in the Thracian basin examine the genotype×environment interaction across the three principal sub-zones of Edirne, Lüleburgaz, and Tekirdağ, and document substantial differences in biotic stress response, most notably the differential expression of yellow rust (Puccinia striiformis), brown rust (Puccinia triticina), and Septoria leaf blotch, across these locations (Öztürk et al., 2018) . Official Trakya-region variety registration trials likewise report yield, quality, powdery mildew, rust, and Septoria tritici observations across Edirne, Lüleburgaz, Keşan, and Tekirdağ locations (Tohumluk Tescil ve Sertifikasyon Merkez Müdürlüğü, 2019) . The Edirne provincial activity report provides a recent administrative surveillance layer for regional pests and plant-health activities, enabling the TRAK-AI knowledge base to distinguish long-standing regional threats from current-season alerts (Edirne İl Tarım ve Orman Müdürlüğü, 2025) ."
        ),
        "For sunflower production, regional research has concentrated": (
            "For sunflower production, regional research has concentrated on three principal threats: downy mildew (Plasmopara halstedii); the parasitic weed broomrape (Orobanche cumana); and field-condition stress symptoms that may be visually confounded with disease. Regional sunflower disease research documents downy-mildew outbreaks in Thrace (Göre, 2009), the races of Plasmopara halstedii detected in Turkish sunflower-growing provinces including Tekirdağ, Edirne, and Kırklareli (Oksal & Maden, 2019), and the effects of broomrape populations collected from Tekirdağ, Kırklareli, and Edirne on sunflower development (Üder & Demirbaş, 2019) . In the domain of remote-sensing-based monitoring, the regional publication on crop-phenology-based object-oriented classification of sunflower parcels in Kırklareli using Sentinel-2 time series provides direct methodological validation for the approach adopted in the present thesis (Karabulut et al., 2021) ."
        ),
        "Studies of irrigation and water-use efficiency are particularly important": (
            "Studies of irrigation and water-use efficiency are particularly important given the climate-adaptation priorities articulated at the national level. Research from Kırklareli and Tekirdağ documents the seasonal water sensitivity of sunflower under regional conditions, and recent Tekirdağ field research characterises the yield response of sunflower under different irrigation timings (Göçmen, 2020) . Edirne-focused studies of supplemental irrigation demonstrate that a single well-timed irrigation event at the start of sunflower flowering produces a measurable yield gain, identifying the most critical irrigation timing for the regional crop calendar (Pekcan & Erdem, 2016) . These primary studies are summarised within the agricultural-assistant subsystem of TRAK-AI DSS and provide the foundation for the irrigation-decision rules implemented in the agronomic-calendar engine described in Chapter 4."
        ),
        "At the soils level, the Thrace soil database project documents": (
            "At the soils level, the Trakya soil database study documents the principal pedological characteristics of the basin, including texture classes, soil reaction, salinity, lime content, organic matter, and macro- and micro-nutrient availability across Edirne, Kırklareli, and Tekirdağ agricultural lands (Gürbüz et al., 2019) . These regional soil attributes are explicitly encoded in the SoilGrids 2.0 layer of the data-fusion pipeline (Poggio et al., 2021) and inform the calibration polynomials applied to the soil-moisture sensors carried by the autonomous rover. Finally, the Edirne İl Tarım ve Orman Müdürlüğü 2025 activity report provides the most recent official surveillance data on disease and pest pressure for the past season, and is included in the knowledge base as a recent-confirmation layer that enables the language model to distinguish between historically common threats and those that have actually exceeded economic thresholds in the current production year (Edirne İl Tarım ve Orman Müdürlüğü, 2025) ."
        ),
        "At the dataset-assembly stage, four academic sources are combined": (
            "At the dataset-assembly stage, four dataset sources are combined to produce a training corpus of approximately twenty-three thousand five hundred images. The Wheat Disease 21K Kaggle dataset contributes 21,212 images of healthy and rust-affected wheat. The Indian Agricultural Research Institute Wheat Nitrogen Deficiency and Rust dataset contributes 859 images that supplement the rust class and introduce the stress_nutrient class. The Bangladesh Agricultural Research Institute Sunflower Disease dataset contributes 1,060 images covering healthy sunflower and downy mildew. A drought-stress image set contributes 360 images for the stress_drought class. The combined corpus is heavily imbalanced, with the healthy_wheat class dominating, and this imbalance is partially addressed by class-weighted loss but remains a known limitation, particularly for the stress_drought class whose 360-image training set raises structural overfit concerns documented in the overfit-warning diagnostic record and revisited in Chapter 5."
        ),
        "At the prompt-construction stage, a structured prompt": (
            "At the prompt-construction stage, a structured prompt is assembled comprising: a system instruction in Turkish that defines the assistant's persona, role, and constraints; a context block enumerating the retrieved passages with explicit source labels; a quantitative-evidence block summarising the current sensor readings, the seven-day NDVI prediction, and the SHAP attribution; and the original user query. The system instruction explicitly directs the model to base its answer on the retrieved passages, to cite source identifiers in its output, and to refuse questions whose answer cannot be located in the knowledge base. The complete prompt typically comprises between 1,000 and 1,500 tokens; on the available CPU hardware, the resulting end-to-end inference latency is approximately 27.1 s. The prompt-engineering procedure follows the principles articulated by (Liu et al., 2022) on in-context learning for retrieval-augmented systems."
        ),
        "The calibration script in calibration/kalibrasyon.py implements": (
            "The calibration script in calibration/kalibrasyon.py implements the polynomial-regression procedure described in Section 3.2. The protocol presents the sensor with ten reference samples spanning the range from oven-dry sand to fully saturated water, recording the raw ten-bit analogue-to-digital-converter reading at each sample. A second-degree polynomial is fit by scipy.optimize.curve_fit, and the resulting coefficients are emitted as a snippet of C++ source code that is then pasted into config.h. The procedure achieves a root mean squared error below 0.03 m³/m³ and a coefficient of determination above 0.95 in the test sequence, consistent with the calibration findings reported by (Abdelmoneim et al., 2025) and the foundational calibration protocols of (Nagahage et al., 2019) . The procedure is documented as a step-by-step recipe in the project's technical documentation, enabling the calibration to be repeated whenever a new sensor is installed or whenever the rover is deployed at a soil profile materially different from the development reference."
        ),
        "The phenological windows used in Figure 4.2 are not arbitrary": (
            "The phenological windows used in Figure 4.2 are not arbitrary: they are derived from the regional research surveyed in Section 2.8, including the Kırklareli sunflower phenology classification studies based on Sentinel-2 time series (Karabulut et al., 2021), the Edirne and Trakya-region winter-wheat trial data (Öztürk et al., 2018; Tohumluk Tescil ve Sertifikasyon Merkez Müdürlüğü, 2019), and the BBCH reference scale (Meier, 2018) . The decision to encode these windows as ranges of calendar months, rather than as exact date pairs, reflects the inter-annual variability documented in the same regional literature: a wheat crop sown on the same date in two consecutive years may diverge in its heading date by ten to fifteen days as a function of accumulated growing-degree-days, and the engine corrects for this variability dynamically by comparing the inferred stage from NDVI dynamics against the expected window."
        ),
        "The SoilGrids 2.0 point sample at the pilot coordinates returns": (
            "The SoilGrids 2.0 point sample at the pilot coordinates returns the following soil profile at the 0–5 cm depth: clay content 30.97 %, sand content 34.99 %, silt content 34.04 %, pH in water 7.11, organic carbon stock 3.42 kg/m², cation exchange capacity 24.3 cmolc/kg, and bulk density 1.28 g/cm³. The textural classification is therefore clay-loam, and the chemical profile is slightly alkaline, both findings consistent with the regional soil-database documentation surveyed in Section 2.8 (Gürbüz et al., 2019) . The 5–15 cm depth exhibits a slight increase in clay content and a corresponding decrease in sand, consistent with the typical pedogenetic process of clay illuviation in temperate soils."
        ),
        "Figure 5.9. Composition of the YOLOv8 training corpus": (
            "Figure 5.9. Composition of the YOLOv8 training corpus. Four datasets contribute the bulk of the corpus: the Wheat Disease 21K Kaggle dataset (21,212 images), the IARI Wheat Nitrogen Deficiency and Rust dataset (859 images), the BARI Sunflower Disease dataset (1,060 images), and the drought-stress image set (360 images). The pronounced imbalance, driven by the dominance of wheat in the Kaggle source, is the principal underlying reason for the per-class accuracy disparities discussed in this section."
        ),
        "The PlantVillage dataset was acquired under controlled laboratory backgrounds": (
            "The PlantVillage dataset was acquired under controlled laboratory backgrounds, with leaves photographed against uniform plain backgrounds at consistent illumination. Subsequent literature established that classifiers trained on PlantVillage routinely lose substantial accuracy when evaluated on field-acquired imagery: (Ferentinos, 2018) and (Barbedo, 2018) both document accuracy losses on the order of fifteen to thirty percentage points when models trained on PlantVillage-like data are evaluated under realistic field conditions. The present 94.9 % accuracy, achieved on a corpus that combines four image-data sources of varying acquisition discipline, including the field-acquired Wheat Disease 21K Kaggle dataset and the IARI Wheat Nitrogen Deficiency and Rust dataset, is therefore plausibly stronger than the 99 % PlantVillage figure when both are evaluated against the same generalisation criterion."
        ),
        "Two per-class results documented in Section 5.4": (
            "Two per-class results documented in Section 5.4, the stress_nutrient class at 85.2 % accuracy and the stress_drought class at a nominal 100 % accuracy from a 360-image training corpus, warrant explicit comparative treatment in this discussion. The stress_nutrient under-performance is consistent with the broader observation in the field-condition image-classification literature that subtle, gradually-presenting stresses are harder to classify from canopy-level imagery than the more visually distinct disease classes (Picon et al., 2019) . The conventional mitigation in the literature is the integration of multispectral or hyperspectral imagery, which captures the chlorophyll-fluorescence and nitrogen-content signals that are invisible in the visible-band imagery available to a consumer-grade camera. This integration is identified as a future-work item for TRAK-AI DSS in Chapter 8."
        ),
    }

    replacements_applied = 0
    for paragraph in doc.paragraphs:
        text = normalized(paragraph.text)
        if text in replacements:
            set_para_text(paragraph, replacements[text])
            replacements_applied += 1
            continue
        if text in exact_reference_replacements:
            set_para_text(paragraph, exact_reference_replacements[text])
            replacements_applied += 1
            continue
        for start, new_text in startswith_replacements.items():
            if text.startswith(start):
                set_para_text(paragraph, new_text)
                replacements_applied += 1
                break

    # Short citation-year fixes in paragraphs that do not otherwise need rewriting.
    short_replacements = {
        "(Picon et al., 2022)": "(Picon et al., 2019)",
        "(Liu et al., 2023)": "(Liu et al., 2022)",
        "(Friedrich, 2017)": "(Friedrich, 2015)",
        "(Edirne Provincial Directorate, 2025)": "(Edirne İl Tarım ve Orman Müdürlüğü, 2025)",
        "(Edirne Provincial Directorate of Agriculture and Forestry, 2025)": "(Edirne İl Tarım ve Orman Müdürlüğü, 2025)",
        "(Turkish Soil Science Society, 2023)": "(Gürbüz et al., 2019)",
        "(Schaefer & Lamm, 2017)": "(Hardie, 2020)",
        "(BARI, 2022)": "",
        "(IARI, 2021)": "",
        "(Kaggle Wheat Disease, 2022)": "",
        "(Yao et al., 2021)": "",
    }
    for paragraph in doc.paragraphs:
        original = paragraph.text
        updated = original
        for old, new in short_replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            set_para_text(paragraph, updated)
            replacements_applied += 1

    delete_exact = {
        "You, J., Li, X., Low, M., Lobell, D., & Ermon, S. (2017). Deep Gaussian process for crop yield prediction based on remote sensing data. In Proceedings of the AAAI Conference on Artificial Intelligence, 31(1), 4559–4565.",
        "Schmitt, M., & Zhu, X. X. (2016). Data fusion and remote sensing: An ever-growing relationship. IEEE Geoscience and Remote Sensing Magazine, 4(4), 6–23.",
        "Wang, C.-Y., Bochkovskiy, A., & Liao, H.-Y. M. (2023). YOLOv7: Trainable bag-of-freebies sets new state-of-the-art for real-time object detectors. In Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR) (pp. 7464–7475).",
        "Turkish Soil Science Society. (2023). Thrace soil database project: Pedological characteristics and nutrient availability mapping. Turkish Soil Science Society Technical Publications.",
        "Yao, J., Tran, S. N., Sawyer, S., & Garg, S. (2021). Drought stress identification in cereal crops using deep learning on small image datasets. Computers and Electronics in Agriculture, 192, 106650.",
    }
    for paragraph in list(doc.paragraphs):
        if paragraph._element is not None and normalized(paragraph.text) in delete_exact:
            delete_paragraph(paragraph)

    doc.save(OUTPUT)
    print(f"saved={OUTPUT}")
    print(f"replacements_applied={replacements_applied}")


if __name__ == "__main__":
    main()
